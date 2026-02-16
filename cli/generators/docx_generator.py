"""DOCX resume generator using python-docx."""

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from ..utils.config import Config
from ..utils.yaml_parser import ResumeYAML


class DocxGenerator:
    """Generate ATS-friendly DOCX resumes."""

    # Standard fonts for ATS compatibility
    FONT_NAME = "Arial"
    FONT_SIZE = 11
    FONT_SIZE_LARGE = 14
    FONT_SIZE_MEDIUM = 12

    def __init__(
        self,
        yaml_path: Optional[Path] = None,
        config: Optional[Config] = None,
    ):
        """
        Initialize DOCX generator.

        Args:
            yaml_path: Path to resume.yaml
            config: Configuration object
        """
        self.yaml_handler = ResumeYAML(yaml_path)
        self.config = config or Config()

    def generate(
        self,
        variant: str,
        output_path: Optional[Path] = None,
        enhanced_context: Optional[Dict[str, Any]] = None,
    ) -> Document:
        """
        Generate resume as DOCX document.

        Args:
            variant: Variant name (e.g., "v1.0.0-backend")
            output_path: Optional output file path
            enhanced_context: Optional dict with AI-enhanced data

        Returns:
            Document object
        """
        # Determine summary key
        variant_key = variant.replace("v1.", "").replace("v2.", "").split("-")[0]
        if variant_key.endswith(".0"):
            variant_key = variant_key.split(".")[0] + "." + variant_key.split(".")[1] + ".0"

        variant_config = self.yaml_handler.get_variant(variant)
        if not variant_config:
            if "backend" in variant:
                summary_key = "backend"
            elif "ml" in variant or "ai" in variant:
                summary_key = "ml_ai"
            elif "fullstack" in variant or "full" in variant:
                summary_key = "fullstack"
            elif "devops" in variant:
                summary_key = "devops"
            elif "leadership" in variant:
                summary_key = "leadership"
            else:
                summary_key = "base"
        else:
            summary_key = variant_config.get("summary_key", "base")

        # Extract technologies for prioritization
        prioritize_technologies = None
        if enhanced_context:
            enhanced_projects = enhanced_context.get("projects", {}).get("featured", [])
            if enhanced_projects:
                all_techs = set()
                for proj in enhanced_projects:
                    if proj.get("highlighted_technologies"):
                        all_techs.update(proj["highlighted_technologies"])
                if all_techs:
                    prioritize_technologies = list(all_techs)

        # Get resume data
        contact = self.yaml_handler.get_contact()
        summary = self.yaml_handler.get_summary(summary_key)
        skills = self.yaml_handler.get_skills(
            variant, prioritize_technologies=prioritize_technologies
        )
        experience = self.yaml_handler.get_experience(variant)
        education = self.yaml_handler.get_education(variant)
        publications = self.yaml_handler.data.get("publications", [])
        certifications = self.yaml_handler.data.get("certifications", [])

        # Merge enhanced context if provided
        if enhanced_context:
            if "summary" in enhanced_context:
                summary = enhanced_context["summary"]
            if "projects" in enhanced_context:
                # Handle enhanced projects structure
                projects = enhanced_context["projects"]
                if isinstance(projects, dict) and "featured" in projects:
                    # Convert to expected format
                    projects_data = {"ai_ml": projects["featured"]}
                else:
                    projects_data = projects
            else:
                projects_data = self.yaml_handler.get_projects(variant)
        else:
            projects_data = self.yaml_handler.get_projects(variant)

        # Create document
        doc = Document()

        # Set default font for entire document
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.FONT_NAME
        font.size = Pt(self.FONT_SIZE)

        # Add content
        self._add_header(doc, contact)
        self._add_summary(doc, summary)
        self._add_projects(doc, projects_data)
        self._add_experience(doc, experience)
        self._add_education(doc, education)
        self._add_skills(doc, skills)
        self._add_publications(doc, publications)
        self._add_certifications(doc, certifications)

        # Save if path provided
        if output_path:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

        return doc

    def _add_header(self, doc: Document, contact: Dict) -> None:
        """Add contact information header."""
        # Name
        name = contact.get("name", "")
        if name:
            para = doc.add_paragraph()
            run = para.add_run(name)
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE_LARGE)
            run.font.bold = True

        # Credentials
        credentials = contact.get("credentials", [])
        if credentials:
            cred_text = ", ".join(credentials)
            para = doc.add_paragraph(cred_text)

        # Location
        location_parts = []
        location = contact.get("location", {})
        if location.get("city"):
            location_parts.append(location["city"])
        if location.get("state"):
            location_parts.append(location["state"])
        if location.get("zip"):
            location_parts.append(location["zip"])

        # Contact info line
        contact_line = []
        if location_parts:
            contact_line.append(", ".join(location_parts))
        if contact.get("email"):
            contact_line.append(contact.get("email"))
        if contact.get("phone"):
            contact_line.append(contact.get("phone"))

        if contact_line:
            para = doc.add_paragraph(" | ".join(contact_line))

        # URLs
        urls = []
        urls_dict = contact.get("urls", {})
        if urls_dict.get("github"):
            urls.append(urls_dict["github"])
        if urls_dict.get("linkedin"):
            urls.append(urls_dict["linkedin"])
        if urls_dict.get("website"):
            urls.append(urls_dict["website"])

        if urls:
            para = doc.add_paragraph(" | ".join(urls))

        doc.add_paragraph()

    def _add_section_heading(self, doc: Document, title: str) -> None:
        """Add a section heading."""
        para = doc.add_paragraph()
        run = para.add_run(title.upper())
        run.font.name = self.FONT_NAME
        run.font.size = Pt(self.FONT_SIZE_MEDIUM)
        run.font.bold = True
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

    def _add_summary(self, doc: Document, summary: Optional[Dict]) -> None:
        """Add professional summary."""
        if not summary:
            return

        self._add_section_heading(doc, "PROFESSIONAL SUMMARY")

        summary_text = summary.get("content", "") if isinstance(summary, dict) else str(summary)
        if summary_text:
            para = doc.add_paragraph(summary_text)
            para.paragraph_format.space_after = Pt(12)

    def _add_projects(self, doc: Document, projects: Dict) -> None:
        """Add projects section."""
        if not projects:
            return

        self._add_section_heading(doc, "PROJECTS")

        for category, project_list in projects.items():
            # Category heading (e.g., "ai_ml" -> "AI/ML")
            category_title = category.replace("_", " ").title()
            if category_title.lower() in ["ai ml", "ai_ml"]:
                category_title = "AI/ML"

            para = doc.add_paragraph()
            run = para.add_run(f"\n{category_title}")
            run.font.bold = True

            for project in project_list:
                # Project name
                project_name = project.get("name", "")
                if project_name:
                    para = doc.add_paragraph()
                    run = para.add_run(project_name)
                    run.font.bold = True

                # Description
                description = project.get("enhanced_description") or project.get("description", "")
                if description:
                    para = doc.add_paragraph(description)
                    para.paragraph_format.left_indent = Inches(0.25)

                # URL
                url = project.get("url")
                if url:
                    para = doc.add_paragraph(f"URL: {url}")
                    para.paragraph_format.left_indent = Inches(0.25)

                # Technologies
                techs = project.get("highlighted_technologies", [])
                if techs:
                    para = doc.add_paragraph(f"Technologies: {', '.join(techs)}")
                    para.paragraph_format.left_indent = Inches(0.25)

                # Achievements
                achievements = project.get("achievement_highlights", [])
                for achievement in achievements:
                    para = doc.add_paragraph(f"- {achievement}")
                    para.paragraph_format.left_indent = Inches(0.5)

        doc.add_paragraph()

    def _add_experience(self, doc: Document, experience: list) -> None:
        """Add work experience section."""
        if not experience:
            return

        self._add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

        for job in experience:
            # Job title and company
            title = job.get("title", "")
            company = job.get("company", "")
            location = job.get("location", "")

            header_text = title
            if company:
                header_text += f" | {company}"
            if location:
                header_text += f" | {location}"

            if header_text:
                para = doc.add_paragraph()
                run = para.add_run(header_text)
                run.font.bold = True

            # Dates
            start_date = job.get("start_date", "")
            end_date = job.get("end_date", "")
            if start_date:
                date_text = f"{start_date} - {end_date if end_date else 'Present'}"
                para = doc.add_paragraph(date_text)
                para.paragraph_format.space_after = Pt(6)

            # Bullets
            bullets = job.get("bullets", [])
            for bullet in bullets:
                bullet_text = bullet.get("text", "") if isinstance(bullet, dict) else str(bullet)
                if bullet_text:
                    para = doc.add_paragraph(f"- {bullet_text}")
                    para.paragraph_format.left_indent = Inches(0.25)

            doc.add_paragraph()

    def _add_education(self, doc: Document, education: list) -> None:
        """Add education section."""
        if not education:
            return

        self._add_section_heading(doc, "EDUCATION")

        for edu in education:
            # Institution and location
            institution = edu.get("institution", "")
            location = edu.get("location", "")

            header_text = institution
            if location:
                header_text += f" | {location}"

            if header_text:
                para = doc.add_paragraph()
                run = para.add_run(header_text)
                run.font.bold = True

            # Degree
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            graduation_date = edu.get("graduation_date", "")

            degree_text = degree
            if field:
                degree_text += f", {field}"
            if graduation_date:
                degree_text += f" | Graduated {graduation_date}"

            if degree_text:
                para = doc.add_paragraph(degree_text)
                para.paragraph_format.left_indent = Inches(0.25)

        doc.add_paragraph()

    def _add_skills(self, doc: Document, skills: Dict) -> None:
        """Add skills section."""
        if not skills:
            return

        self._add_section_heading(doc, "SKILLS")

        for section_name, skill_list in skills.items():
            # Section header
            section_title = section_name.replace("_", " ").title()
            para = doc.add_paragraph()
            run = para.add_run(f"{section_title}: ")
            run.font.bold = True

            # Skills
            skill_texts = []
            for skill in skill_list:
                if isinstance(skill, dict):
                    skill_text = skill.get("name", "")
                    if skill.get("level"):
                        skill_text += f" ({skill['level']})"
                else:
                    skill_text = str(skill)
                if skill_text:
                    skill_texts.append(skill_text)

            if skill_texts:
                run = para.add_run(", ".join(skill_texts))

        doc.add_paragraph()

    def _add_publications(self, doc: Document, publications: list) -> None:
        """Add publications section."""
        if not publications:
            return

        self._add_section_heading(doc, "PUBLICATIONS")

        for pub in publications:
            pub_type = pub.get("type", "")
            title = pub.get("title", "")
            authors = pub.get("authors", "")
            year = pub.get("year", "")
            journal = pub.get("journal", "")
            conference = pub.get("conference", "")

            # Format publication entry
            parts = []
            if authors:
                parts.append(authors)
            if year:
                parts.append(f"({year})")
            if title:
                parts.append(f'"{title}."')

            if journal:
                parts.append(journal)
            elif conference:
                parts.append(f"In {conference}")

            if parts:
                para = doc.add_paragraph("- " + " ".join(parts))

        doc.add_paragraph()

    def _add_certifications(self, doc: Document, certifications: list) -> None:
        """Add certifications section."""
        if not certifications:
            return

        self._add_section_heading(doc, "CERTIFICATIONS")

        for cert in certifications:
            cert_name = cert.get("name", "")
            issuer = cert.get("issuer", "")
            license_num = cert.get("license_number", "")

            text = cert_name
            if issuer:
                text += f" - {issuer}"
            if license_num:
                text += f" (License: {license_num})"

            if text:
                para = doc.add_paragraph(f"- {text}")

        doc.add_paragraph()
